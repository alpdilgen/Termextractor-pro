"""
File parser for TermExtractor-Pro
Supports: TXT, DOCX, PDF, HTML, XML, XLIFF formats
"""

from pathlib import Path
from typing import Optional, Dict, Any
import mimetypes


class FileParser:
    """Parse various file formats"""
    
    SUPPORTED_FORMATS = {
        '.txt': 'text/plain',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.pdf': 'application/pdf',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.xml': 'application/xml',
        '.xliff': 'application/x-xliff',
        '.sdlxliff': 'application/x-sdlxliff',
        '.mqxliff': 'application/x-mqxliff',
    }
    
    def __init__(self):
        """Initialize parser"""
        pass
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse file and extract text.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with 'text' and 'metadata' keys
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                'text': '',
                'metadata': {'error': f'File not found: {file_path}'}
            }
        
        ext = path.suffix.lower()
        
        try:
            if ext == '.txt':
                return self._parse_txt(file_path)
            elif ext == '.docx':
                return self._parse_docx(file_path)
            elif ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.html', '.htm']:
                return self._parse_html(file_path)
            elif ext in ['.xml', '.xliff', '.sdlxliff', '.mqxliff']:
                return self._parse_xml(file_path)
            else:
                # Try as text fallback
                return self._parse_txt(file_path)
        
        except Exception as e:
            return {
                'text': '',
                'metadata': {'error': f'Error parsing file: {str(e)}'}
            }
    
    @staticmethod
    def _parse_txt(file_path: str) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                text = f.read()
            
            return {
                'text': text,
                'metadata': {
                    'format': 'txt',
                    'size_bytes': len(text.encode()),
                    'line_count': len(text.split('\n')),
                }
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                
                return {
                    'text': text,
                    'metadata': {
                        'format': 'txt',
                        'encoding': 'latin-1',
                        'size_bytes': len(text.encode()),
                    }
                }
            except Exception as e:
                return {
                    'text': '',
                    'metadata': {'error': f'Error reading text file: {str(e)}'}
                }
    
    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = '\n\n'.join([para.text for para in doc.paragraphs])
            
            return {
                'text': text,
                'metadata': {
                    'format': 'docx',
                    'paragraph_count': len(doc.paragraphs),
                    'size_bytes': len(text.encode()),
                }
            }
        except ImportError:
            return {
                'text': '',
                'metadata': {'error': 'python-docx not installed. Install with: pip install python-docx'}
            }
        except Exception as e:
            return {
                'text': '',
                'metadata': {'error': f'Error parsing DOCX: {str(e)}'}
            }
    
    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF file"""
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text.append(page.extract_text())
            
            full_text = '\n'.join(text)
            
            return {
                'text': full_text,
                'metadata': {
                    'format': 'pdf',
                    'page_count': len(reader.pages),
                    'size_bytes': len(full_text.encode()),
                }
            }
        except ImportError:
            return {
                'text': '',
                'metadata': {'error': 'PyPDF2 not installed. Install with: pip install PyPDF2'}
            }
        except Exception as e:
            return {
                'text': '',
                'metadata': {'error': f'Error parsing PDF: {str(e)}'}
            }
    
    @staticmethod
    def _parse_html(file_path: str) -> Dict[str, Any]:
        """Parse HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            
            return {
                'text': text,
                'metadata': {
                    'format': 'html',
                    'size_bytes': len(text.encode()),
                }
            }
        except ImportError:
            return {
                'text': '',
                'metadata': {'error': 'BeautifulSoup4 not installed. Install with: pip install beautifulsoup4'}
            }
        except Exception as e:
            return {
                'text': '',
                'metadata': {'error': f'Error parsing HTML: {str(e)}'}
            }
    
    @staticmethod
    def _parse_xml(file_path: str) -> Dict[str, Any]:
        """Parse XML/XLIFF file"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract all text content
            text_parts = []
            
            def extract_text(elem):
                if elem.text:
                    text_parts.append(elem.text)
                for child in elem:
                    extract_text(child)
                    if child.tail:
                        text_parts.append(child.tail)
            
            extract_text(root)
            text = '\n'.join(text_parts)
            
            return {
                'text': text,
                'metadata': {
                    'format': Path(file_path).suffix[1:],
                    'is_bilingual': '<trans-unit' in open(file_path).read()[:5000],
                    'size_bytes': len(text.encode()),
                }
            }
        except Exception as e:
            return {
                'text': '',
                'metadata': {'error': f'Error parsing XML: {str(e)}'}
            }


class ParsedFile:
    """Represents a parsed file"""
    
    def __init__(
        self,
        text: str,
        format: str,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        """Initialize parsed file"""
        self.text = text
        self.format = format
        self.filename = filename
        self.metadata = metadata or {}
    
    def get_word_count(self) -> int:
        """Get approximate word count"""
        return len(self.text.split())
    
    def get_size_bytes(self) -> int:
        """Get size in bytes"""
        return len(self.text.encode())
